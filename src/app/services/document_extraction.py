"""Document text extraction service for various file formats."""

import io
import xml.etree.ElementTree as ET
from typing import Optional

import fitz  # PyMuPDF
from docx import Document

from src.app.common.logging import get_logger

logger = get_logger(__name__)


class DocumentTextExtractor:
    """
    Extract text from various document formats.

    Supports:
    - PDF (using PyMuPDF/fitz)
    - DOCX (using python-docx)
    - TXT (plain text)

    Handles extraction errors gracefully and provides detailed logging.
    """

    # Maximum file size: 50MB
    MAX_FILE_SIZE = 50 * 1024 * 1024

    def extract_text(
        self, content: bytes, content_type: str, file_name: Optional[str] = None
    ) -> str:
        """
        Route to appropriate extractor based on content type.

        Args:
            content: Document content as bytes
            content_type: MIME type (e.g., 'application/pdf', 'text/plain')
            file_name: Optional file name for logging/debugging

        Returns:
            Extracted text as string

        Raises:
            ValueError: If content type is unsupported or content is too large
            Exception: If extraction fails
        """
        # Check file size
        if len(content) > self.MAX_FILE_SIZE:
            raise ValueError(
                f"File too large: {len(content)} bytes (max: {self.MAX_FILE_SIZE} bytes)"
            )

        if len(content) == 0:
            raise ValueError("File is empty (0 bytes)")

        file_info = f" ({file_name})" if file_name else ""
        logger.info(
            f"Extracting text from {content_type}{file_info} - "
            f"size: {len(content)} bytes"
        )

        try:
            # Route to appropriate extractor
            if content_type == "application/pdf":
                text = self._extract_from_pdf(content, file_name)
            elif content_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",
            ]:
                text = self._extract_from_docx(content, file_name)
            elif content_type == "text/plain":
                text = self._extract_from_txt(content, file_name)
            elif content_type in ["text/xml", "application/xml"]:
                text = self._extract_from_xml(content, file_name)
            else:
                # Try to infer from file name if available
                if file_name:
                    inferred_type = self._infer_type_from_filename(file_name)
                    if inferred_type != content_type:
                        logger.warning(
                            f"Content type mismatch - declared: {content_type}, "
                            f"inferred: {inferred_type}. Using inferred type."
                        )
                        return self.extract_text(content, inferred_type, file_name)

                raise ValueError(f"Unsupported content type: {content_type}")

            # Validate extracted text
            if not text or not text.strip():
                logger.warning(f"Extracted text is empty{file_info}")
                return ""

            text_length = len(text)
            logger.info(f"Successfully extracted {text_length} characters{file_info}")

            return text.strip()

        except ValueError:
            # Re-raise ValueError (validation errors)
            raise
        except Exception as e:
            logger.error(
                f"Failed to extract text from {content_type}{file_info}: {str(e)}",
                exc_info=True,
            )
            raise Exception(f"Text extraction failed: {str(e)}") from e

    def _extract_from_pdf(self, content: bytes, file_name: Optional[str] = None) -> str:
        """
        Extract text from PDF using PyMuPDF (fitz).

        Args:
            content: PDF content as bytes
            file_name: Optional file name for logging

        Returns:
            Extracted text

        Raises:
            Exception: If PDF extraction fails
        """
        file_info = f" ({file_name})" if file_name else ""

        try:
            # Open PDF from bytes
            doc = fitz.open(stream=content, filetype="pdf")

            # Extract text from all pages
            text_parts = []
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page.get_text()
                if page_text:
                    text_parts.append(page_text)

            doc.close()

            extracted_text = "\n\n".join(text_parts)

            logger.debug(
                f"Extracted text from PDF{file_info} - "
                f"pages: {len(text_parts)}, chars: {len(extracted_text)}"
            )

            return extracted_text

        except Exception as e:
            logger.error(f"PDF extraction failed{file_info}: {str(e)}", exc_info=True)
            raise Exception(f"Failed to extract text from PDF: {str(e)}") from e

    def _extract_from_docx(
        self, content: bytes, file_name: Optional[str] = None
    ) -> str:
        """
        Extract text from DOCX using python-docx.

        Args:
            content: DOCX content as bytes
            file_name: Optional file name for logging

        Returns:
            Extracted text

        Raises:
            Exception: If DOCX extraction fails
        """
        file_info = f" ({file_name})" if file_name else ""

        try:
            # Open DOCX from bytes
            doc = Document(io.BytesIO(content))

            # Extract text from all paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text_parts.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text_parts.append(cell.text)

            extracted_text = "\n\n".join(text_parts)

            logger.debug(
                f"Extracted text from DOCX{file_info} - "
                f"paragraphs: {len(doc.paragraphs)}, "
                f"tables: {len(doc.tables)}, "
                f"chars: {len(extracted_text)}"
            )

            return extracted_text

        except Exception as e:
            logger.error(f"DOCX extraction failed{file_info}: {str(e)}", exc_info=True)
            raise Exception(f"Failed to extract text from DOCX: {str(e)}") from e

    def _extract_from_txt(self, content: bytes, file_name: Optional[str] = None) -> str:
        """
        Extract text from plain text file.

        Args:
            content: Text content as bytes
            file_name: Optional file name for logging

        Returns:
            Decoded text

        Raises:
            Exception: If decoding fails
        """
        file_info = f" ({file_name})" if file_name else ""

        try:
            # Try UTF-8 first
            try:
                text = content.decode("utf-8")
            except UnicodeDecodeError:
                # Fallback to latin-1 (works for most text files)
                logger.warning(f"UTF-8 decoding failed{file_info}, trying latin-1")
                text = content.decode("latin-1")

            logger.debug(f"Extracted text from TXT{file_info} - chars: {len(text)}")

            return text

        except Exception as e:
            logger.error(f"TXT extraction failed{file_info}: {str(e)}", exc_info=True)
            raise Exception(f"Failed to decode text file: {str(e)}") from e

    def _extract_from_xml(self, content: bytes, file_name: Optional[str] = None) -> str:
        """Extract text from XML/CDA clinical documents by collecting all text nodes."""
        file_info = f" ({file_name})" if file_name else ""

        try:
            root = ET.fromstring(content)

            text_parts = []
            for elem in root.iter():
                if elem.text and elem.text.strip():
                    text_parts.append(elem.text.strip())
                if elem.tail and elem.tail.strip():
                    text_parts.append(elem.tail.strip())

            extracted_text = "\n".join(text_parts)
            logger.debug(f"Extracted text from XML{file_info} - chars: {len(extracted_text)}")
            return extracted_text

        except ET.ParseError as e:
            # Fall back to raw decode if XML is malformed
            logger.warning(f"XML parse error{file_info}: {e} - falling back to raw text decode")
            return self._extract_from_txt(content, file_name)
        except Exception as e:
            logger.error(f"XML extraction failed{file_info}: {str(e)}", exc_info=True)
            raise Exception(f"Failed to extract text from XML: {str(e)}") from e

    def _infer_type_from_filename(self, file_name: str) -> str:
        """
        Infer MIME type from file extension.

        Args:
            file_name: File name with extension

        Returns:
            MIME type
        """
        file_name_lower = file_name.lower()

        if file_name_lower.endswith(".pdf"):
            return "application/pdf"
        elif file_name_lower.endswith(".docx"):
            return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif file_name_lower.endswith(".doc"):
            return "application/msword"
        elif file_name_lower.endswith(".txt"):
            return "text/plain"
        elif file_name_lower.endswith(".xml"):
            return "text/xml"
        else:
            return "application/octet-stream"
